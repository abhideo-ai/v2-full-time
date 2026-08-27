#!/usr/bin/env python3
"""The master résumé as a Word document, rendered from `jobs_tracker_v2`.

    JOBS_TRACKER_DSN=dbname=jobs_tracker_v2   (default)

⚑ .docx IS THE CANONICAL OUTPUT NOW. His words, 2026-08-27: "for resumes, we can
create docx. right? wouldn't that be easier?" … "we create .docx file(s) that I
can also edit easily." … "that's going to be our form now." … "No need of a HTML
page or anything."

    resume_documents + resume_roles + resume_education + resume_profile
    + resume_sections + resume_blocks                          ← AUTHORED
                |
                +--> resume_docx.py generate --> master/Abhisheik_Deo_Resume.docx
                |                                        --> he edits it in Word
                +--> resume_db.py generate --> master/upgrad_resume.generated.html
                                                  ^ ROUND-TRIP PROOF, not a deliverable

This module is a SIBLING of `resume_db.py`, not a replacement. `resume_db.py`
still renders the HTML that proves the database reproduces the master file
byte-for-byte; that gate is what makes trusting these tables safe. This module
renders the same tables to the thing he actually sends. Neither writes
`master/upgrad_resume.html`.

⛔ WHY EVERY WORD COMES OUT OF THE DATABASE.
CLAUDE.md: "Generate copy blocks FROM the résumé, never retype them." Retyping is
how six copies of his career diverged. Nothing here composes a sentence: bullets,
skills, headline, summary, role titles, dates, locations, education and contact
details are all SELECTed. The one string this module supplies is CANDIDATE_NAME
below, and it is flagged there because it is identity, not a claim.

⛔ WHY THE ORDERING IS `resume_sections.ord`, NEVER `section_id`.
`resume_db.render()` selects blocks `ORDER BY section_id, ord`, which is
alphabetical. That is safe THERE only because it immediately regroups the rows
into a dict and lets the ordered `sections` list drive emission. Copying that
query and iterating it directly opens the résumé with the five card-only pre-2016
roles and buries the headline at position 7 — and every count-based check still
passes, because nothing is missing. Order by `s.ord, b.ord` or the document is
wrong in a way that looks right.

⛔ WHY `retired_at IS NULL` SITS IN THE JOIN, NOT THE `WHERE`.
"A bullet LEAVES the master by being retired, never by being deleted"
(migration 012). There are zero retired bullets today, so forgetting the filter
passes now and corrupts later — and `resume_db.load_from()` re-inserts carried
retired bullets at `max(ord)+1` INSIDE their own section, so a stale bullet lands
at the tail of the correct role reading like a real one. In a `WHERE` clause the
filter also degrades the LEFT JOIN to an inner join, and a section whose blocks
were all retired would lose its heading, company and dates entirely.

⛔ WHY BOLD IS THE PASS/FAIL CONDITION.
`resume-issues-to-avoid/README.md` rule 9: bold silently stripped on export makes
the ATS read every bullet as having no highlighted fact and "score craters even
though the source HTML is hygienically correct" — and it was not noticed until an
exported PDF. `<strong>` is not decoration here; it is the bolded fact the
hygiene rules require on every bullet. So `generate` REFUSES to write a document
whose bold spans do not match the database exactly, and `verify` re-opens the
saved file and proves it again from disk.

    resume_docx.py generate [--out P]   render the DB to a .docx
    resume_docx.py verify   [--out P]   re-open that .docx and prove it from the DB

Sibling of `jobs_db.py` / `resume_db.py`: same connection pattern, raw SQL, dict
rows, no ORM. python-docx builds the file and does nothing else.
"""
from __future__ import annotations

import argparse
import hashlib
import os
import re
import sys
import pathlib
from pathlib import Path

import psycopg
from psycopg.rows import dict_row

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

DSN = os.environ.get("JOBS_TRACKER_DSN", "dbname=jobs_tracker_v2")
ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "master" / "Abhisheik_Deo_Resume.docx"
DOC_KEY = "master"

# ⚑ THE ONE STRING NOT IN THE DATABASE. `resume_documents` has no name column and
# master/upgrad_resume.html never carried one — the Hiration card supplied it.
# A name is identity, not a claim, so it is stated here in the open rather than
# smuggled in, and it is the only literal on the page. Everything else is SELECTed.
CANDIDATE_NAME = "Abhisheik Deo"

# ⛔ The ten ids the résumé cannot be missing. Same list as `resume_db.CONTRACT_IDS`,
# imported so the two can never drift.
sys.path.insert(0, str(Path(__file__).resolve().parent))
from resume_db import CONTRACT_IDS, workspace_resume  # noqa: E402

# ═══════════════════════════════════════════════════════════════════════════
# DESIGN TOKENS — the whole look, in one block. Restyling is editing these.
# ═══════════════════════════════════════════════════════════════════════════
# Measured from `master/Abhisheik_Deo_Resume.SUPERSEDED-2026-08-25.docx`, the
# reference he pointed at on 2026-08-27: "is the format for you to use". Every
# value below was re-verified against that file's XML rather than taken on
# trust — see the two corrections flagged inline.
#
# ⛔ WHAT IS DELIBERATELY *NOT* COPIED FROM THE REFERENCE. That file is a
# PDF→Word conversion and its structure is broken: two layout tables, a 217KB
# inline headshot, ten DrawingML shapes, eight `<w:b w:val="0"/>`, eight
# duplicate numbering definitions, and three paragraphs where a heading is
# fused into the text above it. Those damage ATS parsing, which is the entire
# reason this .docx exists. The tokens are taken; the structure is rebuilt.

# ── page ──────────────────────────────────────────────────────────────────
# The reference's A4 is 12/20 twips over the true value (11906×16838) — pdf-lib
# rounding, 0.008in, invisible. True A4 is used here; either is correct.
PAGE_W, PAGE_H = Twips(11906), Twips(16838)

# ⚠ WHY THE MARGINS ARE THIS TIGHT — 0.26in, against a 0.5-1in convention.
# They are the mechanism that makes a three-page résumé achievable. At 64
# bullets the reference fits three pages ONLY because of these margins
# combined with the 8pt body below; widen either and the page count grows.
# The master now carries 89 bullets — 25 more than the reference — so THE
# THREE-PAGE CAP HAS NEVER BEEN VERIFIED AT THIS LENGTH, and cannot be
# verified on this machine (no LibreOffice, and CLAUDE.md forbids driving a
# browser to check rendering). Treat three pages as the intent, not a fact.
MARGIN_TOP, MARGIN_RIGHT = Twips(367), Twips(390)
MARGIN_BOTTOM, MARGIN_LEFT = Twips(484), Twips(397)

# ── the right tab stop ────────────────────────────────────────────────────
# A role is TWO lines, with dates and location pushed to the right margin:
#
#     Principal Software Architect            Mar '25 - Apr '26   ← band
#     VoltusWave Technologies                         Hyderabad   ← no band
#
# ⚠ THIS REVERSES AN EARLIER DECISION IN THIS FILE, AND THE REASON IT GIVES WAS
# WRONG. The previous comment declined the tab stop because "it would change
# what an ATS sees as the employer field", citing a parse of the reference that
# yields `Deque Software⇥Hyderabad` with no title and no dates. That parse is
# real but it comes from the reference's DEGRADED region: body child 80 is a
# 2-column layout table holding the four oldest roles, with the role titles
# leaked into trailing bullet runs. Only 6 of its 11 entries use the clean
# two-line pattern. `Title⇥Dates` / `Company⇥Location` is a standard résumé
# shape and parses correctly — the objection was an artefact of the conversion,
# not a property of the layout. He asked for the format; it was his to choose.
#
# Derived, never hardcoded: the reference's 11131 is 11918 − 397 − 390 on ITS
# page. Ours is true A4, so the same arithmetic gives 11119. A literal would
# silently misalign the moment a margin moves.
ROLE_TAB = Twips(11906 - 397 - 390)

# Dates render `Mar '25 - Apr '26` (his choice, 2026-08-27, matching the
# reference). U+0027 straight apostrophe and U+002D hyphen-minus with spaces —
# verified against all 12 date strings in the reference, which contains no
# U+2019 anywhere and uses en dashes only in body copy, never in dates.
DATE_SEP = " - "

# ── colour ────────────────────────────────────────────────────────────────
ACCENT = RGBColor(0x32, 0x8E, 0xF7)      # 59 uses in the reference; the one accent
ACCENT_HEX = "328EF7"                    # w:pBdr / w:shd take a bare hex string
BAND_HEX = "EBF4FE"                      # pale tint behind each role heading
BODY_COLOR = RGBColor(0, 0, 0)

# ── type scale ────────────────────────────────────────────────────────────
BASE_FONT = "Calibri"                    # the only face in the reference
SZ_NAME = Pt(20)                         # w:sz 40
SZ_SECTION = Pt(11)                      # w:sz 22 — section headings and the headline
SZ_ROLE = Pt(9.5)                        # w:sz 19 — role headings
SZ_CONTACT = Pt(9)                       # w:sz 18 — the contact strip only
SZ_BODY = Pt(8)                          # w:sz 16 — body, bullets, skills, education

# ⚠⚠ THE BRIEF SAID BODY IS 9.5pt. IT IS 8pt, AND THIS WAS RE-VERIFIED HERE.
# `grep -c 'w:sz w:val="19"'` makes 9.5pt look like the most common size. It is
# the most common *explicit* one, and it dresses role and education lines only.
# Resolving style inheritance over every text-bearing run in the reference gives
# the real census — 268 runs at 8pt, 50 at 9.5pt, 6 at 11pt, 1 at 20pt — and
# 64 of its 67 bullet paragraphs carry no `w:sz` at all, inheriting Normal's
# `w:sz w:val="16"`. This is CLAUDE.md's "check the measurement before trusting
# a measurement" trap exactly. Building the body at 9.5pt would run ~19% large
# and would not look like his file.
# ⚠ 8pt is small. It is what his reference does, and it is one constant to
# raise — but raising it costs pages, so decide the two together.

# ── spacing (reference values, twips → points) ────────────────────────────
LINE_SPACING = 259 / 240                 # w:line="259" lineRule="auto" ≈ 1.08
SPACE_BODY_AFTER = Pt(2.25)              # w:after="45"  — Normal
SPACE_SECTION_BEFORE = Pt(8)             # air above the section rule
SPACE_SECTION_AFTER = Pt(2.4)            # w:after="48"  — Heading 1 in the reference
SPACE_ROLE_BEFORE, SPACE_ROLE_AFTER = Pt(4), Pt(5)      # w:after="100"
SPACE_GROUP_BEFORE, SPACE_GROUP_AFTER = Pt(4), Pt(2.25)
SPACE_BULLET_AFTER = Pt(1.95)            # w:after="39"
SPACE_HEADLINE_AFTER = Pt(6)

# ── bullets ───────────────────────────────────────────────────────────────
# The reference's own indents are degenerate — numbering `left=66` fighting a
# paragraph `hanging=162` puts the glyph at −96 twips, inside the margin. These
# are the equivalent geometry arrived at cleanly: glyph at the margin, text and
# every wrapped line at 240 twips.
BULLET_TEXT_INDENT, BULLET_HANGING = Twips(240), Twips(168)
BULLET_GLYPH = "•"                       # reference uses Calibri •, not Symbol

# ── section rule ──────────────────────────────────────────────────────────
# The reference draws five full-width 1.5pt rules as DrawingML shapes in empty
# paragraphs (its `w:pBdr` count is zero). Same look, rebuilt as a real top
# border on the heading style: no shapes, no spacer paragraphs, invisible to an
# ATS. `w:sz` on a border is EIGHTHS of a point, so 1.5pt = 12.
RULE_WEIGHT, RULE_SPACE = 12, 3
CONTACT_BORDER_WEIGHT = 6                # 0.75pt box around the contact strip

# ── justification ─────────────────────────────────────────────────────────
# The reference's Normal is `jc=both`, so its summary and bullets render
# justified. Left here deliberately: at 8pt across a 7.7in measure, justified
# text opens visible word-spacing rivers. One constant to flip back.
BODY_ALIGN = WD_ALIGN_PARAGRAPH.LEFT

# ── style map ─────────────────────────────────────────────────────────────
# Built-in style NAMES only, so he can restyle any of these from Word's style
# pane. The reference defines just eight styles and gives its bullets their look
# from numbering plus direct indents, which is not restylable.
S_NAME = "Heading 1"        # Abhisheik Deo
S_HEADLINE = "Subtitle"     # the tagline
S_SECTION = "Heading 2"     # SUMMARY / KEY SKILLS / EXPERIENCE / EDUCATION
S_ROLE = "Heading 3"        # one role heading line
S_GROUP = "Heading 4"       # a Key Skills group label
S_BULLET = "List Bullet"

LINKEDIN_PREFIX = "linkedin.com/in/"
LINKEDIN_URL = "https://www.linkedin.com/in/"

# The named entities the master uses — the same set `resume_db.KNOWN_ENTITIES`
# guards. `html.unescape()` is deliberately NOT used: it resolves the whole HTML5
# reference table plus semicolon-less legacy forms, so a typo like `&emdash;`
# would silently become something instead of failing loudly. Expand `&amp;` LAST
# so `&amp;lt;` cannot turn into a real `<`.
ENTITIES = [("&lt;", "<"), ("&gt;", ">"), ("&mdash;", "—"), ("&ndash;", "–"),
            ("&nbsp;", " "), ("&quot;", '"'), ("&apos;", "'"), ("&amp;", "&")]
RE_ENTITY = re.compile(r"&#?[A-Za-z0-9]+;")
RE_STRONG = re.compile(r"(</?strong>)")
RE_SKILLS_PREFIX = re.compile(r"^Skills\s*[—–-]\s*")


class BuildError(RuntimeError):
    """The data is not the shape this module knows. Never guessed around."""


# ---------------------------------------------------------------------------
# Inline markup — split on tags FIRST, unescape each run SECOND
# ---------------------------------------------------------------------------
# Order matters and is not cosmetic. 10 of the 16 entities in the master sit
# INSIDE a <strong> (`<strong>25&ndash;30%</strong>`), and unescaping the whole
# string first is the injection `resume_db.py`'s docstring names: `&lt;strong&gt;`
# would become a real tag. Split, then unescape the pieces.

def unescape(text: str) -> str:
    for entity, char in ENTITIES:
        text = text.replace(entity, char)
    stray = RE_ENTITY.search(text)
    if stray:
        raise BuildError(f"unknown HTML entity {stray.group(0)!r} in {text!r}")
    return text


def inline_runs(html: str) -> list[tuple[str, bool]]:
    """`<strong>`-aware split into (text, bold) runs.

    Refuses anything that is not plain text or a balanced <strong>. The master's
    91 live blocks contain exactly one construct — `<strong>` — with no
    attributes and no nesting, so an unrecognised tag is a change to the data
    this module has not been told about, not something to unwrap silently.
    """
    runs: list[tuple[str, bool]] = []
    bold = False
    for piece in RE_STRONG.split(html):
        if piece == "<strong>":
            if bold:
                raise BuildError(f"nested <strong> in {html!r}")
            bold = True
            continue
        if piece == "</strong>":
            if not bold:
                raise BuildError(f"unbalanced </strong> in {html!r}")
            bold = False
            continue
        if not piece:
            continue
        if "<" in piece or ">" in piece:
            raise BuildError(f"unexpected markup {piece!r} in {html!r}")
        runs.append((unescape(piece), bold))
    if bold:
        raise BuildError(f"unclosed <strong> in {html!r}")
    return runs


def strong_spans(html: str) -> list[str]:
    """The bolded facts, in order — what must survive into the .docx."""
    return [text for text, bold in inline_runs(html) if bold]


# ---------------------------------------------------------------------------
# fetch — the whole document, in render order, in one place
# ---------------------------------------------------------------------------

def fetch(doc_key: str = DOC_KEY) -> dict:
    """Every row for one document. `doc_key` defaults to the master; a per-seat
    workspace résumé lives under 'seat:<slug>', loaded by `resume_db.py load --slug`.
    The tables are namespaced by doc_key throughout, so a seat build reads no
    master row and writes none."""
    with psycopg.connect(DSN, row_factory=dict_row) as conn, conn.cursor() as cur:
        cur.execute("SELECT * FROM resume_documents WHERE doc_key = %s", (doc_key,))
        doc = cur.fetchone()
        if doc is None:
            raise BuildError(f"no document {doc_key!r} in {DSN} — run `resume_db.py load`")

        cur.execute("SELECT * FROM resume_profile WHERE doc_key=%s ORDER BY ord", (doc_key,))
        profile = cur.fetchall()
        cur.execute("SELECT * FROM resume_education WHERE doc_key=%s ORDER BY ord", (doc_key,))
        education = cur.fetchall()
        cur.execute("""SELECT * FROM resume_certifications WHERE doc_key=%s ORDER BY ord""",
                    (doc_key,))
        certifications = cur.fetchall()

        # ⛔ ORDER BY s.ord, b.ord — see the module docstring. And `retired_at IS
        # NULL` in the JOIN's ON clause, never in the WHERE.
        cur.execute("""
            SELECT s.ord AS section_ord, s.section_id, s.section_kind, s.card_only,
                   s.heading_html, s.role_n,
                   r.company, r.role_title, r.date_from, r.date_to, r.location,
                   r.date_to_emphasised,
                   b.ord AS block_ord, b.bullet_key, b.html, b.text
              FROM resume_sections s
              LEFT JOIN resume_roles  r ON r.doc_key = s.doc_key AND r.n = s.role_n
              LEFT JOIN resume_blocks b ON b.doc_key = s.doc_key
                                       AND b.section_id = s.section_id
                                       AND b.retired_at IS NULL
             WHERE s.doc_key = %s
             ORDER BY s.ord, b.ord
        """, (doc_key,))
        rows = cur.fetchall()

    sections: list[dict] = []
    for row in rows:
        if not sections or sections[-1]["section_id"] != row["section_id"]:
            sections.append({
                "ord": row["section_ord"], "section_id": row["section_id"],
                "kind": row["section_kind"], "card_only": row["card_only"],
                "heading": unescape(row["heading_html"]), "role_n": row["role_n"],
                "company": row["company"], "role_title": row["role_title"],
                "date_from": row["date_from"], "date_to": row["date_to"],
                "location": row["location"],
                "date_to_emphasised": row["date_to_emphasised"],
                "blocks": [],
            })
        if row["bullet_key"] is not None:
            sections[-1]["blocks"].append(
                {"ord": row["block_ord"], "key": row["bullet_key"],
                 "html": row["html"], "text": row["text"]})

    # ⛔ The contract, asserted before a byte is built. `resume_db.render()` makes
    # the same assertion for the same reason: an empty section is skipped
    # silently downstream and the export still reports success.
    by_id = {s["section_id"]: s for s in sections}
    broken = [i for i in CONTRACT_IDS if not by_id.get(i, {}).get("blocks")]
    if broken:
        raise BuildError(f"parser-contract sections missing or empty: {broken}")

    return {"doc": doc, "profile": profile, "education": education,
            "certifications": certifications, "sections": sections}


def contact(profile: list[dict]) -> list[tuple[str, bool, str | None]]:
    """The contact line as (text, bold, url) parts, from `resume_profile`.

    ⚠ The LinkedIn row's `value_text` is `abhisheikdeo — not abhideo, which is
    the email domain`. That trailing clause is a note to the reader of the HTML
    build artefact, not contact data, and it must never reach a résumé he sends.
    The slug alone is the row's single <strong>, so it is taken from there and
    the gloss cannot follow it.
    """
    fields = {row["field_label"]: row for row in profile}
    for label in ("Current location", "Phone", "Email", "LinkedIn slug"):
        if label not in fields:
            raise BuildError(f"resume_profile has no {label!r} row")

    # Pulled with a regex rather than through `inline_runs`, which refuses
    # anything but <strong>: this one row also carries a <code> tag, and that is
    # correct — the gloss around the slug is prose about the file, and the strict
    # parser refusing it is exactly the signal that it is not résumé content.
    slugs = re.findall(r"<strong>([^<]*)</strong>", fields["LinkedIn slug"]["value_html"])
    if len(slugs) != 1:
        raise BuildError(f"expected exactly one bolded LinkedIn slug, got {slugs!r}")
    slug = unescape(slugs[0])

    parts: list[tuple[str, bool, str | None]] = []
    for text, bold in inline_runs(fields["Current location"]["value_html"]):
        parts.append((text, bold, None))
    parts.append((" · " + fields["Phone"]["value_text"], False, None))
    parts.append((" · " + fields["Email"]["value_text"], False,
                  "mailto:" + fields["Email"]["value_text"]))
    parts.append((" · " + LINKEDIN_PREFIX, False, LINKEDIN_URL + slug))
    parts.append((slug, True, LINKEDIN_URL + slug))
    return parts


# ---------------------------------------------------------------------------
# build
# ---------------------------------------------------------------------------

def _pin_font(style, name: str) -> None:
    """Remove the inherited THEME font reference and pin a real face.

    `style.font.name = "Calibri"` leaves `w:asciiTheme` in place beside the
    explicit `w:ascii`, so the face Word actually uses depends on the theme. It
    is latent rather than active today (this template's major font IS Calibri),
    and it bites the moment anyone picks a different one. Pinning `w:eastAsia`
    and `w:cs` too stops Word substituting a fallback face for the stray glyphs
    this résumé is full of — the em and en dashes, the `·` separators, and the
    single Greek gamma in the Rosenbaum bullet, which sits inside a BOLD run
    where a silent substitution reads as "the bold looks wrong".
    """
    rf = style.element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if qn(attr) in rf.attrib:
            del rf.attrib[qn(attr)]
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(attr), name)


# ⛔ WHY THESE TWO HELPERS EXIST AND WHY THEY INSERT RATHER THAN APPEND.
# python-docx models neither `w:pBdr` (the section rule) nor `w:shd` (the band
# behind a role heading), so both are built by hand. `w:pPr`'s children are a
# SEQUENCE, not a set: pBdr is #9, shd is #10, tabs is #11, spacing #14, ind
# #15, jc #18. Appending `w:shd` to a pPr that already carries a tab stop emits
# it after `w:tabs`, which is schema-invalid — Word then either silently
# repairs the file or declares it corrupt, and neither surfaces until he opens
# it. `insert_element_before` puts each element at its correct position
# whatever order the calls arrive in.
_AFTER_PBDR = ("w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
               "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
               "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
               "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
               "w:textDirection", "w:textAlignment", "w:textboxTightWrap",
               "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr",
               "w:pPrChange")
_AFTER_SHD = _AFTER_PBDR[1:]


def _p_pr(target):
    """The `w:pPr` of a Paragraph or of a ParagraphStyle — they differ.

    ⚠ `getattr(target, "_p", None) or target.element` is WRONG here: an empty
    `<w:p>` is falsy in lxml, so a freshly-created paragraph silently routes to
    the style branch and the formatting lands on the wrong object. Test against
    None explicitly.
    """
    element = getattr(target, "_p", None)
    if element is None:
        element = target.element
    return element.get_or_add_pPr()


def _border(target, edges: tuple[str, ...], weight: int, color: str, space: int = 0) -> None:
    """`weight` is EIGHTHS of a point — 12 is 1.5pt, 6 is 0.75pt. `space` is points."""
    p_pr = _p_pr(target)
    bdr = p_pr.find(qn("w:pBdr"))
    if bdr is None:
        bdr = OxmlElement("w:pBdr")
        p_pr.insert_element_before(bdr, *_AFTER_PBDR)
    for edge in edges:                       # emitted in schema order: top, left, bottom, right
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(weight))
        node.set(qn("w:space"), str(space))
        node.set(qn("w:color"), color)
        bdr.append(node)


def _shading(target, fill: str) -> None:
    """Paragraph-level shading — a full-width band.

    The reference does this with 35 RUN-level `w:shd`, which paints only behind
    the glyphs and breaks the band wherever a run boundary falls. Paragraph
    level is what its own source PDF actually renders, and it is one element
    instead of three per line.
    """
    p_pr = _p_pr(target)
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.insert_element_before(shd, *_AFTER_SHD)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _scrub_stock_subtitle(style) -> None:
    """Strip the two things python-docx's stock `Subtitle` carries that this
    résumé must not inherit.

    Found by unzipping the output and reading styles.xml, not by reasoning about
    it — the stock style ships `<w:numPr><w:ilvl w:val="1"/></w:numPr>`, which
    puts the headline on a list level and indents it, and `<w:spacing
    w:val="15"/>`, which is 0.75pt of letter-spacing. The reference has ZERO
    letter-spacing anywhere in the document, and the headline is not a list item.
    Neither is reachable through python-docx's font/paragraph_format API, and
    neither would have been visible until he opened the file.
    """
    p_pr = style.element.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)
    r_pr = style.element.get_or_add_rPr()
    letter_spacing = r_pr.find(qn("w:spacing"))   # in rPr this is TRACKING, not leading
    if letter_spacing is not None:
        r_pr.remove(letter_spacing)


def _restyle_bullet_glyph(document: Document) -> None:
    """python-docx's `List Bullet` numbering is Symbol U+F0B7, not Calibri `•`.

    Symbol is a legacy symbol-encoded face; where it is unavailable the glyph
    renders as a box or drops entirely, and some PDF text extractors read it as
    a private-use codepoint that lands in the extracted text. The reference pins
    Calibri `•` in all eight of its numbering definitions. One definition is
    enough — the eight are byte-identical apart from indent.

    Removing `w:hint="default"` matters: it is what pulls Symbol back in even
    after the `w:ascii` face has been overridden.
    """
    # Resolve the ONE definition List Bullet actually uses — style → numId →
    # abstractNumId — rather than rewriting every ilvl-0 in the part. The
    # template ships nine numbering definitions and four of them are bulleted;
    # blanket-editing them would silently restyle List Number too.
    style_p_pr = document.styles[S_BULLET].element.find(qn("w:pPr"))
    num_pr = None if style_p_pr is None else style_p_pr.find(qn("w:numPr"))
    if num_pr is None:
        raise BuildError(f"style {S_BULLET!r} carries no numbering reference")
    num_id = num_pr.find(qn("w:numId")).get(qn("w:val"))

    numbering = document.part.numbering_part.element
    num = next((n for n in numbering.findall(qn("w:num"))
                if n.get(qn("w:numId")) == num_id), None)
    if num is None:
        raise BuildError(f"numbering has no w:num with numId {num_id!r}")
    abstract_id = num.find(qn("w:abstractNumId")).get(qn("w:val"))
    abstract = next((a for a in numbering.findall(qn("w:abstractNum"))
                     if a.get(qn("w:abstractNumId")) == abstract_id), None)
    if abstract is None:
        raise BuildError(f"numbering has no abstractNum {abstract_id!r}")

    lvl = next((l for l in abstract.findall(qn("w:lvl"))
                if l.get(qn("w:ilvl")) == "0"), None)
    if lvl is None:
        raise BuildError(f"abstractNum {abstract_id!r} has no ilvl 0")
    lvl.find(qn("w:lvlText")).set(qn("w:val"), BULLET_GLYPH)
    r_pr = lvl.find(qn("w:rPr"))
    fonts = None if r_pr is None else r_pr.find(qn("w:rFonts"))
    if fonts is None:
        raise BuildError(f"abstractNum {abstract_id!r} ilvl 0 has no w:rFonts to repoint")
    if qn("w:hint") in fonts.attrib:
        del fonts.attrib[qn("w:hint")]
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attr), BASE_FONT)

    # Align the numbering's own indent and its `num` tab with the style's, so
    # the two cannot disagree. The style's `w:ind` does win in Word, but the
    # stock definition's left=360/hanging=360 and its num tab at 360 sit 120
    # twips outside BULLET_TEXT_INDENT — the exact fight that leaves the
    # reference's glyph at −96 twips, inside the margin. One source of truth.
    lvl_p_pr = lvl.find(qn("w:pPr"))
    if lvl_p_pr is not None:
        ind = lvl_p_pr.find(qn("w:ind"))
        if ind is not None:
            ind.set(qn("w:left"), str(BULLET_TEXT_INDENT.twips))
            ind.set(qn("w:hanging"), str(BULLET_HANGING.twips))
        tabs = lvl_p_pr.find(qn("w:tabs"))
        tab = None if tabs is None else tabs.find(qn("w:tab"))
        if tab is not None:
            tab.set(qn("w:pos"), str(BULLET_TEXT_INDENT.twips))


def _drop_contextual_spacing(style) -> None:
    """`List Bullet` ships with `<w:contextualSpacing/>`, which makes any spacing
    set between consecutive bullets dead on arrival. python-docx does not expose
    it, so it comes out via lxml or the setting below is a no-op that reads as
    if it worked."""
    p_pr = style.element.get_or_add_pPr()
    found = p_pr.find(qn("w:contextualSpacing"))
    if found is not None:
        p_pr.remove(found)


def _style_document(document: Document) -> None:
    """Every visual decision, applied at STYLE level and nowhere else.

    ⛔ WHY STYLE LEVEL AND NOT DIRECT FORMATTING. Direct run formatting outranks
    a style, so a document formatted run-by-run cannot be restyled from Word's
    style pane at all — the pane appears to do nothing. He edits this file in
    Word; that pane is his only lever. The two things still set per-run are
    bold (which is DATA, one `<strong>` from the database) and colour on the
    non-company half of a role heading (which a style cannot express).
    """
    section = document.sections[0]
    section.page_width, section.page_height = PAGE_W, PAGE_H
    section.top_margin, section.bottom_margin = MARGIN_TOP, MARGIN_BOTTOM
    section.left_margin, section.right_margin = MARGIN_LEFT, MARGIN_RIGHT

    styles = document.styles

    # ── Normal — body, summary, contact, education ────────────────────────
    normal = styles["Normal"]
    normal.font.size = SZ_BODY
    normal.font.color.rgb = BODY_COLOR
    _pin_font(normal, BASE_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = SPACE_BODY_AFTER
    normal.paragraph_format.alignment = BODY_ALIGN
    # ⛔ python-docx's default template carries <w:spacing w:line="276"> in
    # docDefaults — 276/240 = 1.15 line spacing — and NOTHING here overrode it.
    # It applied to all 112 paragraphs, cost roughly half a page, and silently
    # defeated the deliberate spacing set below. Half a page matters: the
    # 3-page cap has never been verified at 89 bullets. The reference's own
    # value is 259/240, which is what LINE_SPACING carries.
    normal.paragraph_format.line_spacing = LINE_SPACING

    # ── the name ─────────────────────────────────────────────────────────
    name = styles[S_NAME]
    name.font.size, name.font.bold, name.font.color.rgb = SZ_NAME, True, ACCENT
    _pin_font(name, BASE_FONT)
    name.paragraph_format.space_before, name.paragraph_format.space_after = Pt(0), Pt(0)

    # ── the headline / tagline ───────────────────────────────────────────
    # ⛔ Bold lives on the STYLE, never on the runs. The headline is the one
    # block the database legitimately stores with no <strong> in it, so bolding
    # its runs would put `<w:b/>` elements in the document that no database
    # value accounts for — inventing emphasis, and muddying the one gate that
    # proves bold survived. A style-level `w:b` renders identically and leaves
    # the run stream a faithful image of the data.
    headline = styles[S_HEADLINE]
    headline.font.size, headline.font.bold = SZ_SECTION, True
    headline.font.color.rgb = ACCENT
    # ⛔ None, NOT False. `= False` writes `<w:i w:val="0"/>`; `= None` removes
    # the stock `<w:i/>` so the style simply inherits Normal. Same reasoning as
    # bold: an explicit negation is a value that has to win an argument, and
    # arguments in Word's style inheritance are what make a style pane useless.
    headline.font.italic = None
    _scrub_stock_subtitle(headline)
    _pin_font(headline, BASE_FONT)
    headline.paragraph_format.space_before = Pt(0)
    headline.paragraph_format.space_after = SPACE_HEADLINE_AFTER
    headline.paragraph_format.alignment = BODY_ALIGN

    # ── section headings, each under its own rule ────────────────────────
    heading = styles[S_SECTION]
    heading.font.size, heading.font.bold = SZ_SECTION, True
    heading.font.color.rgb = ACCENT
    heading.font.all_caps = True   # displays SUMMARY; the extracted text stays "Summary"
    _pin_font(heading, BASE_FONT)
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = SPACE_SECTION_BEFORE
    heading.paragraph_format.space_after = SPACE_SECTION_AFTER
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.keep_together = True
    # The full-width rule. In the reference this is a DrawingML shape sitting in
    # an empty paragraph above the heading; here it is the heading's own top
    # border, which needs no shape, no spacer paragraph and no image part.
    _border(heading, ("top",), RULE_WEIGHT, ACCENT_HEX, RULE_SPACE)

    # ── a role heading ───────────────────────────────────────────────────
    # ⛔ NOT bold at style level, and that is the whole trick. The reference
    # makes its company line a bold style and then needs `<w:b w:val="0"/>` on
    # the location to undo it — eight times. Direct formatting beats the style,
    # so those negations are what make a style pane useless. Leave the style
    # non-bold and let the company run carry the one `<w:b/>` the database asks
    # for, and the same look costs zero negations.
    role = styles[S_ROLE]
    role.font.size, role.font.color.rgb = SZ_ROLE, ACCENT
    role.font.bold = None      # removes the stock <w:b/>; inherits Normal, not bold
    _pin_font(role, BASE_FONT)
    role.paragraph_format.space_before = SPACE_ROLE_BEFORE
    role.paragraph_format.space_after = SPACE_ROLE_AFTER
    role.paragraph_format.keep_with_next = True
    role.paragraph_format.keep_together = True

    # ── a Key Skills group label ─────────────────────────────────────────
    # Deliberately the quietest thing on the page — 8pt, black, not bold — so it
    # reads as a divider between skill bullets rather than competing with a role
    # heading. That is the reference's own treatment, and it is why this needs a
    # style of its own rather than sharing the role heading's.
    group = styles[S_GROUP]
    group.font.size, group.font.color.rgb = SZ_BODY, BODY_COLOR
    group.font.bold = group.font.italic = None    # inherit Normal, never negate
    _pin_font(group, BASE_FONT)
    group.paragraph_format.space_before = SPACE_GROUP_BEFORE
    group.paragraph_format.space_after = SPACE_GROUP_AFTER
    group.paragraph_format.keep_with_next = True

    # ── bullets ──────────────────────────────────────────────────────────
    bullet = styles[S_BULLET]
    bullet.font.size = SZ_BODY
    bullet.font.color.rgb = BODY_COLOR
    _pin_font(bullet, BASE_FONT)
    _drop_contextual_spacing(bullet)
    bullet.paragraph_format.left_indent = BULLET_TEXT_INDENT
    bullet.paragraph_format.first_line_indent = -BULLET_HANGING
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = SPACE_BULLET_AFTER
    bullet.paragraph_format.line_spacing = LINE_SPACING
    bullet.paragraph_format.alignment = BODY_ALIGN
    _restyle_bullet_glyph(document)

    # There is no built-in Hyperlink character style in python-docx's template.
    # Coloured with the accent rather than a second blue: the reference's link
    # colour (0645AD) is an orphan used nowhere else and clashes with the
    # 328EF7 text it sits beside on the contact strip.
    if "Hyperlink" not in [s.name for s in styles]:
        from docx.enum.style import WD_STYLE_TYPE
        link = styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
        link.font.color.rgb = ACCENT
        link.font.underline = True


# A run whose text is this sentinel becomes a <w:tab/> rather than literal text.
# Chosen because it cannot occur in résumé content: `resume_roles` and
# `resume_blocks` are checked for it in check().
TAB = "\x00TAB\x00"

_DATE_RE = re.compile(r"^([A-Z][a-z]{2})\s+(\d{4})$")


def _short_date(value: str) -> str:
    """`Mar 2025` -> `Mar '25`, matching the reference's twelve date strings.

    ⛔ PASSES ANYTHING ELSE THROUGH UNCHANGED. `resume_roles.date_from/date_to`
    are free-text columns, not dates — a value this does not recognise is
    returned as-is rather than mangled, because a silently corrupted date on a
    résumé is far worse than an unconverted one. check() asserts every role's
    dates actually matched, so a pass-through cannot hide.
    """
    v = value.strip()
    m = _DATE_RE.match(v)
    if m:
        return f"{m.group(1)} '{m.group(2)[-2:]}"
    # A RANGE — `resume_education.date_range` holds "Aug 2004 – Dec 2007" as one
    # string, where the role columns hold two. Shorten both halves and rejoin
    # with DATE_SEP, so education reads the same as experience. Without this the
    # two education rows rendered in DIFFERENT formats — "Aug 2004 – Dec 2007"
    # beside "Jan '21" — because only the single-date one matched.
    for dash in ("–", "—", "-"):
        if dash in v:
            left, _, right = v.partition(dash)
            l, r = _DATE_RE.match(left.strip()), _DATE_RE.match(right.strip())
            if l and r:
                return (f"{l.group(1)} '{l.group(2)[-2:]}{DATE_SEP}"
                        f"{r.group(1)} '{r.group(2)[-2:]}")
            break
    return v


def _tab_stop(paragraph) -> None:
    """A single right tab stop on the text edge — see ROLE_TAB."""
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        ROLE_TAB, WD_TAB_ALIGNMENT.RIGHT)


def _add_runs(paragraph, runs, color=None, size=None) -> None:
    """Emit (text, bold) — or (text, bold, color) — runs into a paragraph.

    `color` / `size` are the fallback for runs that do not carry their own.
    Colour and size are safe to set per-run: unlike bold they express something
    a style genuinely cannot (two colours on one line), and neither is data.
    """
    for run_spec in runs:
        text, bold = run_spec[0], run_spec[1]
        run_color = run_spec[2] if len(run_spec) > 2 else color
        if text == TAB:
            # The tab lives in its own run so the runs either side keep their
            # own colour and weight, which is how the reference does it.
            paragraph.add_run().add_tab()
            continue
        run = paragraph.add_run(text)
        # ⛔ Set bold ONLY when true. `run.bold = False` writes an explicit
        # <w:b w:val="0"/>, and direct formatting outranks the style — 299 of
        # them meant that editing "List Bullet" in Word's style pane did
        # NOTHING. Leaving it None inherits the style, which is what makes
        # the document editable.
        if bold:
            run.bold = True
        if run_color is not None:
            run.font.color.rgb = run_color
        if size is not None:
            run.font.size = size


def _add_hyperlink(paragraph, url: str, text: str, bold: bool, size=None) -> None:
    """python-docx 1.2 has no hyperlink API — `paragraph.hyperlinks` is read-only.

    Worth the hand-built element: CLAUDE.md's measurement traps record that the
    LinkedIn slug lived only in a PDF link annotation and a text-only grep read a
    perfectly good file as missing it. Here the slug is in BOTH the relationship
    target and the visible run text, so either check finds it.
    """
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    r_pr.append(style)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if size is not None:
        # `w:sz` is HALF-points, and it must follow `w:b` in the rPr sequence.
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size.pt * 2)))
        r_pr.append(sz)
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    node.set(qn("xml:space"), "preserve")
    run.append(node)
    link.append(run)
    paragraph._p.append(link)


def _role_lines(section: dict) -> tuple[list, list]:
    """Two lines per role, matching the reference:

        Principal Software Architect            Mar '25 - Apr '26   <- band
        VoltusWave Technologies                         Hyderabad   <- no band

    NOT from `heading_html`: that column exists to reproduce the HTML build
    artefact byte-for-byte and carries its wording — `— card-only`, and
    `· THE JAVA YEARS` on El Paso. Both are notes to the reader of that file and
    neither belongs on a résumé he sends.

    ⚑ THE TITLE LEADS, not the company. That is the reference's order, and it is
    the stronger one for this seat class: the title is what a screener matches
    against the role they are filling.

    ⚑ THE DATE IS BLACK AND THE TITLE IS ACCENT — a colour change rather than a
    weight change, so the eye finds the role without the line shouting. `None`
    means inherit the style, which for `S_ROLE` is the accent.

    ⚠ `date_to_emphasised` KEEPS ITS BOLD, AND THIS DEPARTS FROM THE REFERENCE
    DELIBERATELY. The reference bolds no date. But CLAUDE.md, master/README.md
    and the export checklist all require VoltusWave to read Apr 2026, and
    migration 012 records why the flag exists: "the bold is the file shouting
    it. A plain-text store drops it silently and it does not surface until an
    exported PDF." Do not "restore fidelity" by removing this.
    """
    dates = _short_date(section["date_from"]) + DATE_SEP
    line1 = [(section["role_title"], True, None),
             (TAB, False, None),
             (dates, False, BODY_COLOR),
             (_short_date(section["date_to"]),
              bool(section["date_to_emphasised"]), BODY_COLOR)]
    line2 = [(section["company"], True, None)]
    if section["location"]:
        line2 += [(TAB, False, None), (section["location"], False, BODY_COLOR)]
    return line1, line2


def build(model: dict) -> Document:
    document = Document()
    _style_document(document)

    document.add_paragraph(CANDIDATE_NAME, style=S_NAME)

    # ── the contact strip ────────────────────────────────────────────────
    # The reference renders this as a bordered, tinted, full-width band — and
    # builds it out of a THREE-CELL LAYOUT TABLE. The band is the design; the
    # table is conversion damage that an ATS reads as three columns. One
    # centred paragraph with a four-sided border and a fill is the same look
    # with a single text flow.
    line = document.add_paragraph(style="Normal")
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_before = Pt(2)
    line.paragraph_format.space_after = Pt(6)
    _border(line, ("top", "left", "bottom", "right"), CONTACT_BORDER_WEIGHT, ACCENT_HEX)
    _shading(line, BAND_HEX)
    for text, bold, url in contact(model["profile"]):
        if url:
            _add_hyperlink(line, url, text, bold, size=SZ_CONTACT)
        else:
            run = line.add_run(text)
            # ⛔ Set bold ONLY when true. `run.bold = False` writes an explicit
            # <w:b w:val="0"/>, and direct formatting outranks the style — 299 of
            # them meant that editing "List Bullet" in Word's style pane did
            # NOTHING. Leaving it None inherits the style, which is what makes
            # the document editable.
            if bold:
                run.bold = True
            run.font.size = SZ_CONTACT
            run.font.color.rgb = ACCENT

    by_id = {s["section_id"]: s for s in model["sections"]}

    # Headline — the only block with no bold, legitimately: the DB's own
    # `resume_blocks_bold` constraint is scoped to experience sections. Do not
    # assert bold here and never invent it. Its accent colour and weight come
    # from the Subtitle style, so no run here carries formatting of its own.
    headline = document.add_paragraph(style=S_HEADLINE)
    _add_runs(headline, inline_runs(by_id["quick-headline"]["blocks"][0]["html"]))

    document.add_paragraph("Summary", style=S_SECTION)
    summary = document.add_paragraph(style="Normal")
    _add_runs(summary, inline_runs(by_id["quick-summary"]["blocks"][0]["html"]))

    document.add_paragraph("Key Skills", style=S_SECTION)
    for section in model["sections"]:
        if section["kind"] != "skills":
            continue
        document.add_paragraph(RE_SKILLS_PREFIX.sub("", section["heading"]),
                               style=S_GROUP)
        for block in section["blocks"]:
            _add_runs(document.add_paragraph(style=S_BULLET),
                      inline_runs(block["html"]))

    document.add_paragraph("Professional Experience", style=S_SECTION)
    for section in model["sections"]:
        if section["kind"] != "experience":
            continue
        # The pale band behind each role heading is the reference's strongest
        # single wayfinding device — it is what lets the eye find the eight role
        # boundaries in a dense three-page document. Paragraph-level shading, so
        # it spans the full width instead of breaking at every run boundary the
        # way the reference's 35 run-level fills do.
        line1, line2 = _role_lines(section)
        # Line 1 carries the band. PARAGRAPH-level shading, deliberately: the
        # reference paints 35 RUN-level fills, which break the band at every run
        # boundary — including across the tab gap, where it matters most. One
        # element instead of three, and a band that cannot break.
        head = document.add_paragraph(style=S_ROLE)
        _tab_stop(head)
        _shading(head, BAND_HEX)
        _add_runs(head, line1)
        # Line 2 is unshaded, which is what separates the two visually.
        sub = document.add_paragraph(style=S_ROLE)
        _tab_stop(sub)
        _add_runs(sub, line2)
        for block in section["blocks"]:
            _add_runs(document.add_paragraph(style=S_BULLET),
                      inline_runs(block["html"]))

    document.add_paragraph("Education", style=S_SECTION)
    for row in model["education"]:
        # His reference builds an education entry EXACTLY like a role heading —
        # shaded band, accent credential, 9.5pt. Rendering it as plain 8pt body
        # copy was the one place this file diverged from his format without a
        # reason, so it is the only defect the verifiers held the promote on.
        # The reference builds education EXACTLY like a role — two lines, band
        # on the first, dates to the right margin (its body children 85/86 and
        # 87/88). Same treatment here, from the same helpers.
        head = document.add_paragraph(style=S_ROLE)
        _tab_stop(head)
        _shading(head, BAND_HEX)
        _add_runs(head, [(row["credential"], True, None),
                         (TAB, False, None),
                         (_short_date(row["date_range"]), False, BODY_COLOR)])
        sub = document.add_paragraph(style=S_ROLE)
        _add_runs(sub, [(row["institution"], True, None)])

    # ⛔ `resume_certifications` is DELIBERATELY EMPTY — migration 012: "His
    # certifications live only in the Hiration card … The emptiness is the record
    # of that boundary", and "inventing rows to make 'the database is the source'
    # sound complete is the shape of the bug migration 007 paid for on
    # source_url." So the heading is emitted only if there is something under it.
    # The one certification-shaped item in the corpus, the IIIT PG Diploma, is
    # stored as education and renders above, which is where the master puts it.
    if model["certifications"]:
        document.add_paragraph("Certifications", style=S_SECTION)
        for row in model["certifications"]:
            para = document.add_paragraph(style="Normal")
            runs = [(row["name"], True)]
            if row["issuer"]:
                runs.append((", " + row["issuer"], False))
            if row["awarded"]:
                runs.append((" (" + row["awarded"] + ")", False))
            _add_runs(para, runs)

    return document


# ---------------------------------------------------------------------------
# check — THE GATE, run against a Document object whatever its provenance
# ---------------------------------------------------------------------------

def _paragraphs(document: Document) -> list[dict]:
    out = []
    for para in document.paragraphs:
        # ⛔ NOT para.runs. python-docx returns only DIRECT-CHILD <w:r>, so every
        # run nested inside a <w:hyperlink> is invisible — the bolded LinkedIn
        # slug among them. That is why the report printed 223 bold spans one line
        # above 224 <w:b/> elements: a reporting artefact that also left bold
        # inside any hyperlink UNGATED. Walk the XML instead, in document order.
        runs = []
        for r in para._p.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"):
            texts = r.findall(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
            text = "".join(t.text or "" for t in texts)
            rpr = r.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
            bold = False
            if rpr is not None:
                b = rpr.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b")
                bold = b is not None and b.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") != "0"
            runs.append((text, bold))
        out.append({"style": para.style.name, "text": para.text, "runs": runs})
    return out


def _bold_spans(runs: list[tuple[str, bool]]) -> list[str]:
    """Contiguous bold runs, joined — what a reader sees as one bolded fact."""
    spans, current = [], []
    for text, bold in runs:
        if bold:
            current.append(text)
        elif current:
            spans.append("".join(current))
            current = []
    if current:
        spans.append("".join(current))
    return spans


def check(document: Document, model: dict) -> list[str]:
    """Every failure, listed. Returns [] when the document is right.

    Never returns early: a report naming one problem when there are four is how
    a second run finds a "new" defect that was there all along.
    """
    problems: list[str] = []
    paras = _paragraphs(document)
    by_style: dict[str, list[dict]] = {}
    for para in paras:
        by_style.setdefault(para["style"], []).append(para)
    bullets = by_style.get("List Bullet", [])

    # 1. Every live block present, text byte-identical to the DB's plain column.
    expected = [b for s in model["sections"] if s["kind"] in ("skills", "experience")
                for b in s["blocks"]]
    got = [p["text"] for p in bullets]
    if len(got) != len(expected):
        problems.append(f"bullet count {len(got)} != {len(expected)} live skills+experience blocks")
    for block, text in zip(expected, got):
        if text != block["text"]:
            problems.append(f"bullet {block['key']} text differs\n"
                            f"      db   : {block['text']!r}\n"
                            f"      docx : {text!r}")

    # 2. Headline and summary, likewise — they are Normal paragraphs, not bullets.
    for section_id in ("quick-headline", "quick-summary"):
        block = next(s for s in model["sections"] if s["section_id"] == section_id)["blocks"][0]
        if block["text"] not in [p["text"] for p in paras]:
            problems.append(f"{section_id} paragraph missing or altered")

    # 3. BOLD — the pass/fail condition. Every <strong> span in the database
    #    becomes a bold span in the document, in order.
    db_spans = [span for s in model["sections"] for b in s["blocks"]
                for span in strong_spans(b["html"])]
    doc_spans = [span for p in paras for span in _bold_spans(p["runs"])]
    # ⛔ MULTIPLICITY, not membership. `s not in doc_spans` asks only whether the
    # text appears bold SOMEWHERE. Of 209 database spans just 182 are distinct —
    # "Java and Spring Boot" occurs four times — so bold could vanish from three
    # of those four and a membership test would still pass. 27 occurrences were
    # invisible to the gate. resume-issues-to-avoid/ rule 9 is precisely this
    # failure (bold silently stripped on export, every bullet then flagged "no
    # bolded fact"), so the one detector for it must not be blind to repeats.
    from collections import Counter
    want, have = Counter(db_spans), Counter(doc_spans)
    short = {k: want[k] - have.get(k, 0) for k in want if want[k] > have.get(k, 0)}
    if short:
        total = sum(short.values())
        problems.append(f"{total} of {len(db_spans)} database <strong> occurrence(s) "
                        f"are NOT bold in the .docx: "
                        f"{[(k, f'want {want[k]}, have {have.get(k, 0)}') for k in list(short)[:5]]}")

    # 4. All fifteen sections reachable, all ten roles headed with their dates.
    #    ⚠ Skills groups and role headings now live in DIFFERENT styles — the
    #    reference treats them as opposites (a role heading is 9.5pt bold accent
    #    on a band; a group label is deliberately the quietest text on the page)
    #    and one style cannot be both. Look each up where it is actually emitted,
    #    via the S_* constants, so a future restyle cannot leave this asserting
    #    against a style nothing is written in — which would pass by finding
    #    nothing to disagree with.
    role_heads = [p["text"] for p in by_style.get(S_ROLE, [])]
    group_heads = [p["text"] for p in by_style.get(S_GROUP, [])]
    for section in model["sections"]:
        if section["kind"] == "skills":
            name = RE_SKILLS_PREFIX.sub("", section["heading"])
            if name not in group_heads:
                problems.append(f"skills group heading missing: {name!r}")
        elif section["kind"] == "experience":
            # A role is TWO paragraphs now. Assert BOTH, and assert they are
            # ADJACENT and in order — a title line orphaned from its company
            # line renders as two unrelated bands and no automated check below
            # would notice. The tab is a real <w:tab/>, which extracts as \t.
            l1, l2 = _role_lines(section)
            want1 = "".join(r[0] for r in l1).replace(TAB, "\t")
            want2 = "".join(r[0] for r in l2).replace(TAB, "\t")
            try:
                i = role_heads.index(want1)
            except ValueError:
                problems.append(f"role line 1 missing: {want1!r}")
                continue
            if i + 1 >= len(role_heads) or role_heads[i + 1] != want2:
                problems.append(f"role line 2 missing or not adjacent to line 1: "
                                f"{want2!r}")
        # ⛔ The dates must have SHORTENED. _short_date passes an unrecognised
        # value through unchanged by design, so without this a schema change or
        # a stray value would silently ship `Mar 2025` beside `Mar '25`.
        if section["kind"] == "experience":
            for col in ("date_from", "date_to"):
                if _short_date(section[col]) == section[col] and section[col]:
                    problems.append(f"date not shortened for {section['company']!r}: "
                                    f"{section[col]!r} — _short_date did not match it")
    for label in ("Summary", "Key Skills", "Professional Experience", "Education"):
        if label not in [p["text"] for p in by_style.get(S_SECTION, [])]:
            problems.append(f"section heading missing: {label!r}")
    if model["certifications"] and "Certifications" not in [
            p["text"] for p in by_style.get(S_SECTION, [])]:
        problems.append("certifications exist in the DB but no Certifications heading")

    # 5. Education, contact, name.
    for row in model["education"]:
        if _short_date(row["date_range"]) == row["date_range"]:
            problems.append(f"education date not shortened: {row['date_range']!r} — "
                            f"_short_date did not match it")
        # Education is two paragraphs now, same as a role: credential + dates,
        # then institution. Assert each separately rather than on one line.
        if not any(row["credential"] in p["text"] for p in paras):
            problems.append(f"education credential missing: {row['credential']!r}")
        if not any(row["institution"] in p["text"] for p in paras):
            problems.append(f"education institution missing: {row['institution']!r}")
    wanted_contact = "".join(t for t, _, _ in contact(model["profile"]))
    if not any(p["text"] == wanted_contact for p in paras):
        problems.append(f"contact line missing or altered — wanted {wanted_contact!r}")
    if not any(p["text"] == CANDIDATE_NAME for p in by_style.get(S_NAME, [])):
        problems.append("name heading missing")

    # 6. ATS layout — the parsing failure `resume-issues-to-avoid/` names.
    body = document.element.body
    if body.findall(".//" + qn("w:tbl")):
        problems.append("document contains a table — ATS parsers mis-read multi-column layout")
    if body.findall(".//" + qn("w:txbxContent")):
        problems.append("document contains a text box")
    for para in paras:
        if "•" in para["text"]:
            problems.append(f"literal bullet glyph in text: {para['text'][:60]!r}")

    # 7. THE RESTYLE'S OWN INVARIANTS — the reference's damage, kept out.
    #    Every one of these is something the reference file actually does, and
    #    the whole point of rebuilding rather than cloning it is that none of
    #    them survives. Asserted here so a later styling change cannot quietly
    #    reintroduce one.
    body_xml = body.xml
    negations = len(re.findall(r'<w:b w:val="0"', body_xml))
    if negations:
        # ⚠ SCOPED TO document.xml ON PURPOSE. python-docx's own template ships
        # seven <w:b w:val="0"/> inside unused MediumGrid2* TABLE styles; a
        # whole-package grep would fail on the template's noise forever and the
        # real gate would get switched off to silence it.
        problems.append(f"{negations} <w:b w:val=\"0\"/> in the body — direct formatting "
                        f"beats the style and makes Word's style pane useless")
    # ⚠ AND ON THE STYLES THIS RÉSUMÉ USES. Setting `style.font.bold = False`
    # writes the same negation into styles.xml, where the body-scoped check
    # above cannot see it — this restyle introduced exactly that on two heading
    # styles and only an unzip caught it. `= None` is the fix: it removes the
    # stock element and lets the style inherit. Scoped to the styles actually
    # used, because the stock template ships seven of these on unused
    # MediumGrid2* TABLE styles and a whole-file count could never reach zero.
    for style_name in (S_NAME, S_HEADLINE, S_SECTION, S_ROLE, S_GROUP, S_BULLET, "Normal"):
        r_pr = document.styles[style_name].element.find(qn("w:rPr"))
        bold = None if r_pr is None else r_pr.find(qn("w:b"))
        if bold is not None and bold.get(qn("w:val")) in ("0", "false"):
            problems.append(f"style {style_name!r} carries <w:b w:val=\"0\"/> — "
                            f"set .font.bold = None to inherit instead of negating")

    if body.findall(".//" + qn("w:drawing")) or body.findall(".//" + qn("w:pict")):
        problems.append("document contains a drawing or picture — the reference's rules, "
                        "icons and headshot are all shapes, and none is ATS-readable")
    # ⚠ NUMBERING LIVES IN THE STYLE, NOT ON THE PARAGRAPHS. `List Bullet`
    # carries `w:numPr` in styles.xml, so document.xml holds none and a grep for
    # numPr over the body reads a perfectly good file as having no bullets —
    # this assertion was written that way first and failed on correct output.
    # Real list numbering is what matters (an ATS reads it as a list, and he can
    # restyle every bullet at once); assert it where it is.
    if bullets:
        style_p_pr = document.styles[S_BULLET].element.find(qn("w:pPr"))
        style_num = None if style_p_pr is None else style_p_pr.find(qn("w:numPr"))
        if style_num is None and not body.findall(".//" + qn("w:numPr")):
            problems.append(f"{len(bullets)} bullet paragraph(s) but neither the "
                            f"{S_BULLET!r} style nor any paragraph carries w:numPr — "
                            f"they are not real list items")

    return problems


def _report(document: Document, model: dict) -> None:
    paras = _paragraphs(document)
    db_spans = [span for s in model["sections"] for b in s["blocks"]
                for span in strong_spans(b["html"])]
    doc_spans = [span for p in paras for span in _bold_spans(p["runs"])]
    xml = document.element.body.xml
    section = document.sections[0]
    print(f"  blocks in DB      : {sum(len(s['blocks']) for s in model['sections'])}")
    print(f"  paragraphs        : {len(paras)}")
    print(f"  bullets           : {sum(1 for p in paras if p['style'] == S_BULLET)}")
    print(f"  role headings     : {sum(1 for p in paras if p['style'] == S_ROLE)}")
    print(f"  skills groups     : {sum(1 for p in paras if p['style'] == S_GROUP)}")
    print(f"  DB <strong> spans : {len(db_spans)}")
    print(f"  bold spans in docx: {len(doc_spans)}")
    print(f"  <w:b/> elements   : {len(re.findall(r'<w:b/>', xml))}")
    print(f"  <w:b w:val=\"0\"/>  : {len(re.findall(r'<w:b w:val=.0.', xml))}  (must be 0)")
    print(f"  xml:space=preserve: {len(re.findall(r'xml:space=\"preserve\"', xml))}")
    print(f"  page · margins    : {section.page_width.twips}×{section.page_height.twips} tw · "
          f"{section.top_margin.twips}/{section.right_margin.twips}/"
          f"{section.bottom_margin.twips}/{section.left_margin.twips} tw")
    print(f"  accent · band     : #{ACCENT_HEX} ×{len(re.findall(ACCENT_HEX, xml))} · "
          f"#{BAND_HEX} ×{len(re.findall(BAND_HEX, xml))}")
    # ⛔ COUNTED OVER THE ELEMENT TREE, NEVER `re.findall(qn(...), xml)`.
    # `qn("w:tbl")` expands to `{http://…/main}tbl`, which is lxml's internal
    # form and appears NOWHERE in serialized XML — the serialized form is
    # `<w:tbl>`. So the old `tables · textboxes` line printed 0 · 0 whatever the
    # document contained: a gate-shaped number that could not fail. (The gate in
    # `check()` was always correct; it uses `body.findall`. Only the report
    # lied.) CLAUDE.md's measurement traps, one more time.
    body = document.element.body

    def count(tag: str) -> int:
        return len(body.findall(".//" + qn(tag)))

    # ⚠ Scoped to the section-heading style, not a whole-file pBdr count: the
    # stock template also puts borders on `Title` and `Intense Quote`, neither
    # of which this résumé uses. A file-wide 3 would read as two rules that do
    # not exist.
    section_p_pr = document.styles[S_SECTION].element.find(qn("w:pPr"))
    has_rule = section_p_pr is not None and section_p_pr.find(qn("w:pBdr")) is not None
    print(f"  tables · textboxes: {count('w:tbl')} · {count('w:txbxContent')}   (must be 0 · 0)")
    print(f"  drawings · picts  : {count('w:drawing')} · {count('w:pict')}   (must be 0 · 0)")
    print(f"  body shd · pBdr   : {count('w:shd')} · {count('w:pBdr')}"
          f"   (10 role bands + contact strip · contact strip)")
    print(f"  section rule      : {'yes' if has_rule else 'NO'} — a top border on the "
          f"{S_SECTION!r} style, not a shape")
    print(f"  {S_BULLET} numPr : "
          f"{'yes' if document.styles[S_BULLET].element.find(qn('w:pPr')) is not None else 'NO'}"
          f" (style-level; document.xml correctly holds none)")


# ---------------------------------------------------------------------------

def generate(out_path: Path, doc_key: str = DOC_KEY) -> int:
    model = fetch(doc_key)
    document = build(model)

    # ⛔ CHECKED BEFORE IT IS WRITTEN. A silently-wrong résumé on disk is worse
    # than no résumé: he opens it, it looks finished, and the defect surfaces in
    # a recruiter's ATS. Same refusal posture as `resume_db.generate`.
    # ⛔ The stock python-docx template ships docProps that say
    # <dc:creator>python-docx</dc:creator>, an empty <dc:title/> and
    # dcterms:created 2013-12-23. All visible in Word's Info pane and in macOS
    # Get Info, and some applicant tracking systems ingest dc:title as the
    # document's name. On a file he uploads to job boards that is a template
    # watermark. Stamp it with his own details instead.
    core = document.core_properties
    core.author = core.last_modified_by = "Abhisheik Deo"
    core.title = "Abhisheik Deo — Resume"
    core.subject = "Principal Software Architect"   # role family, not a per-seat title
    core.comments = ""
    core.category = core.keywords = ""

    problems = check(document, model)
    if problems:
        print(f"[resume-docx] REFUSING to write {out_path} — "
              f"{len(problems)} problem(s):", file=sys.stderr)
        for problem in problems:
            print(f"    ⛔ {problem}", file=sys.stderr)
        return 1

    # ⛔ NEVER over the round-trip verification artefact. `resume_db.py` guards the
    # same file the same way; without this, `--out master/upgrad_resume.html`
    # replaces the one artefact that proves a parse did not drop a <strong> —
    # with a binary zip.
    protected = (ROOT / "master" / "upgrad_resume.html").resolve()
    if out_path.resolve() == protected:
        raise SystemExit(f"[resume-docx] REFUSING: {out_path} is the round-trip "
                         f"verification artefact, not an output path")
    if out_path.suffix.lower() != ".docx":
        raise SystemExit(f"[resume-docx] REFUSING: {out_path} is not a .docx path")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    # Atomic: a crash mid-save must not leave a truncated .docx that Word opens
    # as corrupt. Same pattern as resume_db.generate.
    import tempfile
    with tempfile.NamedTemporaryFile(dir=str(out_path.parent), suffix=".docx",
                                     delete=False) as tmp:
        tmp_path = pathlib.Path(tmp.name)
    document.save(str(tmp_path))
    tmp_path.replace(out_path)
    digest = hashlib.sha256(out_path.read_bytes()).hexdigest()
    print(f"[resume-docx] wrote {out_path} "
          f"({out_path.stat().st_size:,} bytes, sha256 {digest[:8]}…)")
    _report(document, model)
    return 0


def verify(out_path: Path, doc_key: str = DOC_KEY) -> int:
    if not out_path.exists():
        print(f"[resume-docx] no file at {out_path} — run `generate` first", file=sys.stderr)
        return 2
    model = fetch(doc_key)
    document = Document(str(out_path))   # re-read from disk, not the object in memory
    print(f"{'=' * 68}\n  DOCX GATE — {out_path}\n{'=' * 68}")
    print(f"  sha256            : {hashlib.sha256(out_path.read_bytes()).hexdigest()}")
    _report(document, model)
    problems = check(document, model)
    if problems:
        print(f"\n  ⛔ {len(problems)} problem(s):")
        for problem in problems:
            print(f"    - {problem}")
        return 1
    print("\n  ✅ Every live block present and byte-identical to the database, "
          "every <strong> span bold, every section and role heading present.")
    if not model["certifications"]:
        print("  ⚠ resume_certifications holds 0 rows — deliberately (migration 012). "
              "No Certifications section is emitted, and none is invented.")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(
        description="The master résumé as a Word document, rendered from jobs_tracker_v2.")
    sub = parser.add_subparsers(dest="cmd", required=True)
    for name, helptext in (("generate", "render the DB to a .docx"),
                           ("verify", "re-open the .docx and prove it from the DB")):
        p = sub.add_parser(name, help=helptext)
        p.add_argument("--out", type=Path, default=None,
                       help=f"output path (default {OUT}, or the workspace with --slug)")
        p.add_argument("--doc-key", default=None,
                       help="document key (default 'master'); exclusive with --slug")
        p.add_argument("--slug", default=None,
                       help="workspace slug -> doc_key 'seat:<slug>', written into "
                            "that workspace as Abhisheik_Deo_Resume.docx")

    args = parser.parse_args()
    try:
        # ⛔ Both name the document. Preferring one silently is how a seat gets
        # built from the master's rows and nobody notices until he opens it.
        if args.slug and args.doc_key:
            raise SystemExit("[resume-docx] pass --slug OR --doc-key, not both")
        doc_key = args.doc_key or (f"seat:{args.slug}" if args.slug else DOC_KEY)
        out = args.out
        if out is None:
            out = (workspace_resume(args.slug).parent / "Abhisheik_Deo_Resume.docx"
                   if args.slug else OUT)
        # ⛔ A seat build must not land on the master .docx, and a master build
        # must not land in a workspace. Both are silent, plausible-looking wrongs.
        if doc_key != DOC_KEY and out.resolve() == OUT.resolve():
            raise SystemExit(f"[resume-docx] REFUSING: {out} is the MASTER .docx and "
                             f"doc_key is {doc_key!r}")
        if doc_key == DOC_KEY and out.resolve() != OUT.resolve():
            raise SystemExit(f"[resume-docx] REFUSING: writing the master document to "
                             f"{out} — pass --slug or --doc-key if this is a seat")
        return (generate(out, doc_key) if args.cmd == "generate"
                else verify(out, doc_key))
    except BuildError as exc:
        print(f"[resume-docx] REFUSING: {exc}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    sys.exit(main())
